import io

import pytest

from app.ingest.parsers import SUPPORTED_EXTENSIONS, parse_bytes


def test_plain_text_and_markdown():
    assert parse_bytes("notes.txt", b"hello world") == "hello world"
    assert parse_bytes("readme.MD", b"# Title\nbody") == "# Title\nbody"


def test_html_strips_tags_and_scripts():
    html = b"""<html><head><title>skip</title><style>p{color:red}</style></head>
    <body><script>var x = 'skip me';</script><h1>Heading</h1><p>Body text.</p></body></html>"""
    text = parse_bytes("page.html", html)
    assert "Heading" in text
    assert "Body text." in text
    assert "skip me" not in text
    assert "color:red" not in text


def test_csv_rows_joined():
    text = parse_bytes("data.csv", b"name,role\nada,engineer\n")
    assert "name, role" in text
    assert "ada, engineer" in text


def test_tsv_dialect_detected():
    text = parse_bytes("data.tsv", b"name\trole\nada\tengineer\n")
    assert "ada, engineer" in text


def test_docx_paragraphs_and_tables():
    import docx

    document = docx.Document()
    document.add_paragraph("Word document paragraph.")
    table = document.add_table(rows=1, cols=2)
    table.rows[0].cells[0].text = "cell one"
    table.rows[0].cells[1].text = "cell two"
    buf = io.BytesIO()
    document.save(buf)

    text = parse_bytes("report.docx", buf.getvalue())
    assert "Word document paragraph." in text
    assert "cell one, cell two" in text


def test_unsupported_extension_raises():
    with pytest.raises(ValueError, match="Unsupported file type"):
        parse_bytes("archive.zip", b"PK...")
    with pytest.raises(ValueError, match="Unsupported file type"):
        parse_bytes("no_extension", b"data")


def test_pdf_is_a_registered_format():
    assert ".pdf" in SUPPORTED_EXTENSIONS
    assert ".docx" in SUPPORTED_EXTENSIONS
