"""File parsers: turn uploaded bytes into plain text for the chunking pipeline.

Extension-based registry so new formats are one function + one dict entry away.
Heavy parsers (pypdf, python-docx) import lazily so the app starts without them
until a file of that type actually arrives.
"""

from __future__ import annotations

import csv
import io
from html.parser import HTMLParser


def _parse_text(data: bytes) -> str:
    return data.decode("utf-8", errors="ignore")


class _HTMLTextExtractor(HTMLParser):
    _SKIP = {"script", "style", "noscript", "head"}

    def __init__(self) -> None:
        super().__init__()
        self._chunks: list[str] = []
        self._skip_depth = 0

    def handle_starttag(self, tag, attrs):
        if tag in self._SKIP:
            self._skip_depth += 1

    def handle_endtag(self, tag):
        if tag in self._SKIP and self._skip_depth:
            self._skip_depth -= 1

    def handle_data(self, data):
        if not self._skip_depth and data.strip():
            self._chunks.append(data.strip())

    @property
    def text(self) -> str:
        return "\n".join(self._chunks)


def _parse_html(data: bytes) -> str:
    extractor = _HTMLTextExtractor()
    extractor.feed(_parse_text(data))
    return extractor.text


def _parse_csv(data: bytes) -> str:
    text = _parse_text(data)
    dialect = "excel-tab" if "\t" in text.splitlines()[0] else "excel"
    rows = csv.reader(io.StringIO(text), dialect=dialect)
    return "\n".join(", ".join(cell.strip() for cell in row if cell.strip()) for row in rows)


def _parse_pdf(data: bytes) -> str:
    from pypdf import PdfReader

    reader = PdfReader(io.BytesIO(data))
    return "\n".join(page.extract_text() or "" for page in reader.pages)


def _parse_docx(data: bytes) -> str:
    import docx

    document = docx.Document(io.BytesIO(data))
    parts = [p.text for p in document.paragraphs if p.text.strip()]
    for table in document.tables:
        for row in table.rows:
            cells = [c.text.strip() for c in row.cells if c.text.strip()]
            if cells:
                parts.append(", ".join(cells))
    return "\n".join(parts)


_PARSERS = {
    ".txt": _parse_text,
    ".md": _parse_text,
    ".rst": _parse_text,
    ".log": _parse_text,
    ".html": _parse_html,
    ".htm": _parse_html,
    ".csv": _parse_csv,
    ".tsv": _parse_csv,
    ".pdf": _parse_pdf,
    ".docx": _parse_docx,
}

SUPPORTED_EXTENSIONS = frozenset(_PARSERS)


def parse_bytes(filename: str, data: bytes) -> str:
    """Extract plain text from ``data`` based on the file extension.

    Raises ValueError for unsupported extensions (callers map this to HTTP 415).
    """
    dot = filename.rfind(".")
    ext = filename[dot:].lower() if dot != -1 else ""
    parser = _PARSERS.get(ext)
    if parser is None:
        supported = ", ".join(sorted(SUPPORTED_EXTENSIONS))
        raise ValueError(f"Unsupported file type {ext or '(none)'!r}; supported: {supported}")
    return parser(data)
